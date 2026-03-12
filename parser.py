"""
parser.py — Extract plain text from PDF, DOCX, and TXT files.

Dependencies:
  pip install pymupdf python-docx
"""

import io
import re
from typing import Optional


def extract_text(content_bytes: bytes, extension: str) -> Optional[str]:
    """
    Extract plain text from document bytes.

    Args:
        content_bytes: Raw bytes of the document.
        extension:     File extension string (e.g. '.pdf', '.docx', '.txt').

    Returns:
        Extracted text string, or None on failure.
    """
    ext = extension.lower()

    if ext == ".pdf":
        return _extract_pdf(content_bytes)
    elif ext == ".docx":
        return _extract_docx(content_bytes)
    elif ext in (".txt", ".md", ".csv"):
        return _extract_text(content_bytes)
    else:
        # Attempt UTF-8 decoding as a fallback
        try:
            return content_bytes.decode("utf-8", errors="replace")
        except Exception:
            return None


def _extract_pdf(content_bytes: bytes) -> Optional[str]:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=content_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return _clean_text("\n".join(pages))
    except ImportError:
        # Fallback to pdfplumber if PyMuPDF not available
        return _extract_pdf_pdfplumber(content_bytes)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}") from e


def _extract_pdf_pdfplumber(content_bytes: bytes) -> Optional[str]:
    """Fallback PDF extraction using pdfplumber."""
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return _clean_text("\n".join(pages))
    except ImportError:
        raise RuntimeError(
            "No PDF parser available. Install PyMuPDF: pip install pymupdf"
        )
    except Exception as e:
        raise RuntimeError(f"PDF extraction (pdfplumber) failed: {e}") from e


def _extract_docx(content_bytes: bytes) -> Optional[str]:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return _clean_text("\n".join(paragraphs))
    except ImportError:
        raise RuntimeError(
            "python-docx not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}") from e


def _extract_text(content_bytes: bytes) -> Optional[str]:
    """Decode a plain text / TXT / Markdown / CSV file."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return _clean_text(content_bytes.decode(encoding))
        except UnicodeDecodeError:
            continue
    return content_bytes.decode("utf-8", errors="replace")


def _clean_text(text: str) -> str:
    """Normalise whitespace and remove junk characters from extracted text."""
    # Collapse runs of whitespace (preserve single newlines)
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ consecutive newlines → 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()


def truncate_text(text: str, max_chars: int = 12_000) -> str:
    """
    Truncate text to a maximum character count for LLM context limits.
    Tries to break at a sentence boundary.
    """
    if len(text) <= max_chars:
        return text
    truncated = text[:max_chars]
    # Find last sentence boundary
    last_period = max(
        truncated.rfind(". "),
        truncated.rfind(".\n"),
        truncated.rfind("! "),
        truncated.rfind("? "),
    )
    if last_period > max_chars * 0.8:
        return truncated[: last_period + 1] + "\n\n[... document truncated for summarization ...]"
    return truncated + "\n\n[... document truncated for summarization ...]"
